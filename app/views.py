from flask import Blueprint, render_template, request, flash, redirect, url_for, send_file
from .models import Product
from sqlalchemy import or_
import os
from docx import Document
from docxcompose.composer import Composer
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from io import BytesIO


views = Blueprint('views', __name__)


@views.route('/', methods = ['GET', 'POST'])
def home():
    products = Product.query.all()
    categories = set([product.category for product in products])
    return render_template('base.html', categories=categories, products = products)


def replace_text(paragraph, tag, value, apply_styling_func, selected_size):
    if tag in paragraph.text:
        for run in paragraph.runs:
            if tag in run.text:
                new_text = run.text.replace(tag, str(value))
                run.clear()
                run.add_text(new_text)
                apply_styling_func(run, tag, selected_size)

def apply_styling(run, tag, selected_size):
    styles = {
        'A6': {
            '<Name>': {'size': 55, 'bold': True, 'italic': False, 'name': 'Marion'},
            '<D>': {'size': 22, 'bold': True, 'italic': False, 'name': 'Avenir Next LT Pro'},
            '<R>': {'size': 21, 'bold': False, 'italic': False, 'name': 'Avenir LT Std 55 Roman'},
            '<N>': {'size': 82, 'bold': True, 'italic': False, 'name': 'Avenir Next LT Pro'}
        },
        'A7': {
            '<Name>': {'size': 40, 'bold': True, 'italic': False, 'name': 'Marion'},
            '<D>': {'size': 16, 'bold': True, 'italic': False, 'name': 'Avenir Next LT Pro'},
            '<R>': {'size': 16, 'bold': False, 'italic': False, 'name': 'Avenir LT Std 55 Roman'},
            '<N>': {'size': 60, 'bold': True, 'italic': False, 'name': 'Avenir Next LT Pro'}
        },
        'A4': {
            '<Name>': {'size': 110, 'bold': True, 'italic': False, 'name': 'Marion'},
            '<D>': {'size': 28, 'bold': True, 'italic': False, 'name': 'Avenir Next LT Pro'},
            '<R>': {'size': 24, 'bold': False, 'italic': False, 'name': 'Avenir LT Std 55 Roman'},
            '<N>': {'size': 140, 'bold': True, 'italic': False, 'name': 'Avenir Next LT Pro'}
        }
    }

    if selected_size in styles and tag in styles[selected_size]:
        style = styles[selected_size][tag]
        run.font.size = Pt(style['size'])
        run.bold = style['bold']
        run.italic = style['italic']
        run.font.name = style['name']
        run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

@views.route('/submit-products', methods=['POST'])
def submit_products():
    selected_size = request.form.get('selected_size')

    products = []
    for key in request.form.keys():
        if key.startswith('product_discounted_price_'):
            #Extract the unique index and product ID
            parts = key.split('_')
            index = parts[3]
            product_id = parts[4]

            #Creating new product objects with the data from frontend
            product = Product.query.get(product_id)
            product_name = request.form.get(f'product_name_{index}_{product_id}')
            discounted_price = request.form.get(f'product_discounted_price_{index}_{product_id}')

            #Adding the instance with unique details
            products.append({
                'id': product_id,
                'name': product_name,
                'discounted_price': discounted_price,
                'original_price': product.original_price
            })
    
    flash('Prints Ready!', category='success')
    
    template_file = f"{os.getcwd()}/app/templates/{selected_size}.docx"

    main_doc = Document()
    composer = None

    for index, product in enumerate(products, start=1):
        product_name_parts = product['name'].split('|')
        product_name_parts_init = product_name_parts[0].split(' ')

        #Replacing tags in the template document
        document = Document(template_file)
        if index != len(products) and selected_size != 'A7':
            document.add_page_break()

        for paragraph in document.paragraphs:
            replace_text(paragraph, '<Name>', product_name_parts_init[0], apply_styling, selected_size)
            remaining_parts = ' '.join(product_name_parts_init[1:])
            replace_text(paragraph, '<D>', remaining_parts, apply_styling, selected_size)
            replace_text(paragraph, '<R>', product['original_price'], apply_styling, selected_size)
            replace_text(paragraph, '<N>', product['discounted_price'], apply_styling, selected_size)

        if index > 1:
            if not composer:
                composer = Composer(main_doc)
            composer.append(document)
        else:
            main_doc = document

    #Creating a BytesIO object with the document
    byte_io = BytesIO()
    if composer:
        composer.save(byte_io)
    else:
        main_doc.save(byte_io)
    byte_io.seek(0)

    #Downloading file
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    if len(products) > 1:
        filename = f"{timestamp}_multi.docx"
    else:
        filename = f"{product_name_parts_init[0]}.docx"

    return send_file(byte_io, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
